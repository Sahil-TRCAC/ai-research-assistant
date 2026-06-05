"""
services/document_service.py
─────────────────────────────
Business logic for document management.

Data Flow (Upload)
──────────────────
  1. Receive FileStorage from Flask route
  2. Validate extension against ALLOWED_EXTENSIONS
  3. Save file to UPLOAD_FOLDER with UUID-prefixed name
  4. If PDF → call pdf_extractor.extract(file_path)
       ├─ success  → store full text, preview, word/char count
       └─ failure  → store error, mark is_scanned if applicable
  5. Persist Document row to PostgreSQL (status=ready | error)
  6. Return Document instance to the route layer
"""

import os
import uuid
import logging
from datetime import datetime, timezone
from werkzeug.datastructures import FileStorage
from werkzeug.utils import secure_filename
from flask import current_app

from models import db
from models.document import Document
from services.pdf_extractor import extract as extract_pdf

logger = logging.getLogger("ai_research.document_service")

_PREVIEW_CHARS = 500


# ── Private helpers ───────────────────────────────────────────────────────────

def _allowed_extension(filename: str) -> bool:
    allowed = current_app.config.get("ALLOWED_EXTENSIONS", set())
    return "." in filename and filename.rsplit(".", 1)[1].lower() in allowed


def _unique_filename(original: str) -> str:
    """Prepend a UUID hex so stored filenames are always collision-free."""
    ext = original.rsplit(".", 1)[-1].lower() if "." in original else "bin"
    return f"{uuid.uuid4().hex}.{ext}"


def _extract_text_for_file(file_path: str, file_type: str) -> dict:
    """
    Dispatch to the appropriate extractor based on file type.

    Returns a dict with keys matching Document columns:
        extracted_text, content_preview, word_count, char_count,
        page_count, is_scanned, extraction_error
    """
    result = {
        "extracted_text": None,
        "content_preview": None,
        "word_count": None,
        "char_count": None,
        "page_count": None,
        "is_scanned": False,
        "extraction_error": None,
    }

    # ── PDF ──────────────────────────────────────────────────────────────────
    if file_type == "pdf":
        ex = extract_pdf(file_path)
        result["page_count"] = ex.page_count
        result["is_scanned"] = ex.is_scanned
        result["extraction_error"] = ex.error

        if ex.success:
            result["extracted_text"] = ex.full_text
            result["content_preview"] = ex.preview(_PREVIEW_CHARS)
            result["word_count"] = ex.word_count
            result["char_count"] = ex.char_count

    # ── Plain text / Markdown ────────────────────────────────────────────────
    elif file_type in ("txt", "md", "csv"):
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as fh:
                text = fh.read()
            result["extracted_text"] = text
            result["content_preview"] = text[:_PREVIEW_CHARS]
            result["word_count"] = len(text.split())
            result["char_count"] = len(text)
        except Exception as exc:
            logger.warning("Could not read text file %s: %s", file_path, exc)
            result["extraction_error"] = str(exc)

    # ── DOCX ─────────────────────────────────────────────────────────────────
    elif file_type == "docx":
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text = "\n".join(p.text for p in doc.paragraphs)
            result["extracted_text"] = text
            result["content_preview"] = text[:_PREVIEW_CHARS]
            result["word_count"] = len(text.split())
            result["char_count"] = len(text)
            result["page_count"] = len(doc.sections)
        except Exception as exc:
            logger.warning("Could not extract DOCX %s: %s", file_path, exc)
            result["extraction_error"] = str(exc)

    else:
        result["extraction_error"] = f"Text extraction not supported for .{file_type} files."

    return result


# ── Public API ────────────────────────────────────────────────────────────────

def upload_document(
    file: FileStorage,
    title: str | None = None,
    description: str | None = None,
    tags: list | None = None,
) -> Document:
    """
    Save an uploaded file to disk, extract its text, and persist to DB.

    Parameters
    ----------
    file        : Werkzeug FileStorage from the multipart request.
    title       : Human-readable title (defaults to original filename).
    description : Optional free-text description.
    tags        : Optional list of string tags.

    Returns
    -------
    Document instance committed to the database.
    """
    if not file or not file.filename:
        raise ValueError("No file provided.")

    original_filename = secure_filename(file.filename)
    if not _allowed_extension(original_filename):
        raise ValueError(f"File type not allowed: {original_filename!r}")

    # ── 1. Save file to disk ─────────────────────────────────────────────────
    stored_filename = _unique_filename(original_filename)
    upload_dir = current_app.config["UPLOAD_FOLDER"]
    os.makedirs(upload_dir, exist_ok=True)
    file_path = os.path.join(upload_dir, stored_filename)

    file.save(file_path)
    file_size = os.path.getsize(file_path)
    file_type = original_filename.rsplit(".", 1)[-1].lower()
    mime_type = file.mimetype or None

    logger.info(
        "Saved upload: %s -> %s (%d bytes)", original_filename, stored_filename, file_size
    )

    # ── 2. Extract text ──────────────────────────────────────────────────────
    logger.info("Extracting text from: %s (type=%s)", stored_filename, file_type)
    extraction = _extract_text_for_file(file_path, file_type)

    # ── 3. Determine final status ────────────────────────────────────────────
    if extraction["extracted_text"]:
        status = "ready"
    elif extraction["is_scanned"]:
        status = "error"          # scanned PDFs cannot be processed yet
    elif extraction["extraction_error"] and file_type == "pdf":
        status = "error"
    else:
        # Non-PDF or unknown type without extraction support
        status = "ready"

    # ── 4. Persist to PostgreSQL ─────────────────────────────────────────────
    doc = Document(
        title=title or original_filename,
        original_filename=original_filename,
        stored_filename=stored_filename,
        file_path=file_path,
        file_size=file_size,
        file_type=file_type,
        mime_type=mime_type,
        status=status,
        description=description,
        tags=tags or [],
        # Content fields from extraction
        extracted_text=extraction["extracted_text"],
        content_preview=extraction["content_preview"],
        page_count=extraction["page_count"],
        word_count=extraction["word_count"],
        char_count=extraction["char_count"],
        is_scanned=extraction["is_scanned"],
        extraction_error=extraction["extraction_error"],
    )
    db.session.add(doc)
    db.session.commit()

    logger.info(
        "Document persisted: id=%s status=%s words=%s",
        doc.id, doc.status, doc.word_count,
    )

    # ── 5. Auto-chunk if extraction succeeded ────────────────────────────────────
    if doc.extracted_text:
        try:
            from services.chunking_service import chunk_document
            chunks = chunk_document(doc.id)
            logger.info("Auto-chunking complete: doc=%s  chunks=%d", doc.id, len(chunks))
        except Exception as exc:  # noqa: BLE001
            logger.warning("Auto-chunking failed for doc=%s: %s", doc.id, exc)

    return doc


def get_all_documents(
    page: int = 1,
    per_page: int = 20,
    status: str | None = None,
) -> tuple[list[Document], int]:
    """Return a paginated list of documents (newest first)."""
    query = Document.query.order_by(Document.created_at.desc())
    if status:
        query = query.filter_by(status=status)
    total = query.count()
    docs = query.offset((page - 1) * per_page).limit(per_page).all()
    return docs, total


def get_document_by_id(document_id: str) -> Document | None:
    """Return a single Document or None."""
    return Document.query.get(document_id)


def get_document_text(document_id: str) -> Document | None:
    """Return a Document loaded with its full extracted_text (same as get_by_id)."""
    return Document.query.get(document_id)


def delete_document(document_id: str) -> bool:
    """Delete a document record and remove the file from disk."""
    doc = Document.query.get(document_id)
    if not doc:
        return False

    if os.path.exists(doc.file_path):
        try:
            os.remove(doc.file_path)
            logger.info("Deleted file: %s", doc.file_path)
        except OSError as exc:
            logger.warning("Could not remove file %s: %s", doc.file_path, exc)

    db.session.delete(doc)
    db.session.commit()
    logger.info("Document deleted: id=%s", document_id)
    return True


def update_document_metadata(document_id: str, **kwargs) -> Document | None:
    """Update mutable metadata fields (title, description, tags)."""
    doc = Document.query.get(document_id)
    if not doc:
        return None

    allowed_fields = {"title", "description", "tags"}
    for field, value in kwargs.items():
        if field in allowed_fields:
            setattr(doc, field, value)

    doc.updated_at = datetime.now(timezone.utc)
    db.session.commit()
    logger.info("Document updated: id=%s fields=%s", document_id, list(kwargs.keys()))
    return doc


def re_extract_document(document_id: str) -> Document | None:
    """
    Re-run text extraction on an existing document's file.
    Useful for retrying failed or scanned documents.
    """
    doc = Document.query.get(document_id)
    if not doc:
        return None

    if not os.path.exists(doc.file_path):
        doc.status = "error"
        doc.extraction_error = "File not found on disk — cannot re-extract."
        db.session.commit()
        return doc

    logger.info("Re-extracting document: id=%s", document_id)
    extraction = _extract_text_for_file(doc.file_path, doc.file_type)

    doc.extracted_text = extraction["extracted_text"]
    doc.content_preview = extraction["content_preview"]
    doc.word_count = extraction["word_count"]
    doc.char_count = extraction["char_count"]
    doc.page_count = extraction["page_count"]
    doc.is_scanned = extraction["is_scanned"]
    doc.extraction_error = extraction["extraction_error"]
    doc.status = "ready" if extraction["extracted_text"] else "error"
    doc.updated_at = datetime.now(timezone.utc)

    db.session.commit()
    logger.info("Re-extraction done: id=%s status=%s", doc.id, doc.status)

    # Re-chunk if extraction succeeded
    if doc.extracted_text:
        try:
            from services.chunking_service import chunk_document
            chunks = chunk_document(doc.id)
            logger.info("Re-chunking complete: doc=%s  chunks=%d", doc.id, len(chunks))
        except Exception as exc:  # noqa: BLE001
            logger.warning("Re-chunking failed for doc=%s: %s", doc.id, exc)

    return doc
